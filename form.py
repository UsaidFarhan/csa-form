import streamlit as st
import tempfile
import os
import json
import requests
import ast
from docx import Document

# Gemini API key (replace this with your key securely in production!)
GEMINI_API_KEY = st.secrets["API-KEY"]

# Gemini API endpoint
GEMINI_MODEL_NAME = "models/gemini-1.5-flash"  # Flash model
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1/{GEMINI_MODEL_NAME}:generateContent"

# Function to call Gemini API
def call_gemini(prompt):
    headers = {
        "Content-Type": "application/json",
    }
    params = {"key": GEMINI_API_KEY}
    body = {
        "contents": [
            {
                "parts": [
                    {
                        "text": prompt
                    }
                ]
            }
        ]
    }

    response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=body)
    if response.status_code == 200:
        try:
            text = response.json()['candidates'][0]['content']['parts'][0]['text']
            return text.strip()
        except Exception as e:
            return f"Error parsing Gemini response: {e}"
    else:
        return f"Error: {response.status_code}, {response.text}"

# Function to fill CSA form
def fill_csa_form(template_path, output_path, data_dict):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data_dict.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_dict.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

    doc.save(output_path)

# Streamlit app
st.set_page_config(page_title="CSA Form Generator", page_icon="📄")
st.title("📄 Contract Summary Approval (CSA) Generator")

# Upload Contract File
uploaded_file = st.file_uploader("Upload Fully Negotiated Contract (PDF or DOCX)", type=["pdf", "docx"])

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[-1]) as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_file_path = temp_file.name

    st.success(f"Uploaded file: {uploaded_file.name}")

    if st.button("Extract Data from Contract"):
        try:
            contract_text = ""
            if uploaded_file.name.endswith(".docx"):
                from docx import Document
                doc = Document(temp_file_path)
                contract_text = "\n".join([para.text for para in doc.paragraphs])
            elif uploaded_file.name.endswith(".pdf"):
                from PyPDF2 import PdfReader
                reader = PdfReader(temp_file_path)
                contract_text = "\n".join([page.extract_text() for page in reader.pages])

            st.info("Extracting data using Gemini...")

            gemini_prompt = f"""
You are a legal contract summarizer. Extract the following information from the contract text below.

For each field, return:
- The exact text from the contract, if it exists.
- The value null (without quotes) if the field is not present.

Return the result strictly as JSON, starting with '{{' and ending with '}}'. No explanations, no code blocks, no markdown. The field names must exactly match the list below.

Fields to extract:
- Supplier Name: The name of the supplier company.
- Contract Title: The title or subject of the contract.
- Contract Type: Choose exactly one of the following options — "New", "Renewal", or "Addendum". Return exactly one of these options in the output. Do not return any other text.
- Brief Scope of Service/Supply: A short description of the services or goods provided.
- Contract Value: The total value or price in the contract.
- Payment Terms: Any details about payment structure or schedule.
- Delivery Timelines: Dates or timeframes for delivery.
- Warranty Details: Details of any warranty provided.
- Payment Currency: The currency used for payment (e.g., USD, PKR).
- Penalties: Any penalties or charges mentioned for delays or non-performance.
- Date of Contract: The date when the contract was signed or agreed upon.
- Effective Date: The date when the contract terms start to apply.
- Commencement Date of Supply/Service (if different from Effective Date): The start date of actual supply or service, if different from the Effective Date.
- Term or Duration of Service/Supply: The length or duration of the contract.
- Notice Period for Termination: The period required to give notice before terminating the contract.

Contract text:
{contract_text}
"""



            gemini_response = call_gemini(gemini_prompt)
            st.code(gemini_response, language="json")

            # Parse JSON safely
            try:
                try:
                    extracted_json = json.loads(gemini_response)
                except json.JSONDecodeError:
                    extracted_json = ast.literal_eval(gemini_response)

                if isinstance(extracted_json, dict):
                    st.session_state['extracted_dict'] = extracted_json
                    st.success("Data extracted successfully!")
                    st.json(extracted_json)
                else:
                    st.warning("Gemini response is not a JSON object. Please check the output.")
                    st.code(gemini_response)
            except Exception as e:
                st.error(f"Error parsing Gemini response: {e}")
                st.code(gemini_response)

        except Exception as e:
            st.error(f"Error processing file: {e}")

# Generate CSA Form Button
if 'extracted_dict' in st.session_state:
    st.markdown("### Generate CSA Form:")
    csa_template = "F1 Contract Summary Approval Form.docx"  # Replace with your actual template file

    if st.button("Generate CSA Form"):
        output_file = "CSA_Filled.docx"
        try:
            fill_csa_form(csa_template, output_file, st.session_state['extracted_dict'])

            with open(output_file, "rb") as f:
                st.download_button(
                    label="Download CSA Form",
                    data=f,
                    file_name="CSA_Filled.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            st.success("CSA Form generated successfully!")
        except Exception as e:
            st.error(f"Error generating CSA form: {e}")
